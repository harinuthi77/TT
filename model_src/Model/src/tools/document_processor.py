from typing import Dict
import os

class DocumentProcessor:
    def process_pdf(self, filepath: str) -> Dict:
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return {'text': text, 'pages': len(reader.pages)}
        except:
            return {'text': '', 'pages': 0}
    
    def process_docx(self, filepath: str) -> Dict:
        try:
            from docx import Document
            doc = Document(filepath)
            text = '\n'.join([p.text for p in doc.paragraphs])
            return {'text': text, 'paragraphs': len(doc.paragraphs)}
        except:
            return {'text': '', 'paragraphs': 0}
    
    def summarize(self, text: str, max_length: int = 500) -> str:
        words = text.split()
        if len(words) <= max_length:
            return text
        return ' '.join(words[:max_length]) + '...'
